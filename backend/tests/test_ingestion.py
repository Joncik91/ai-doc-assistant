"""Tests for document extraction and chunking."""

from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument

from app.ingestion.chunker import chunk_pages
from app.ingestion.extractors import extract_document
from app.ingestion.extractors import ExtractedPage


def _write_simple_pdf(path: Path, text: str) -> None:
    safe_text = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream = f"BT /F1 24 Tf 72 72 Td ({safe_text}) Tj ET"
    objects = [
        "1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n",
        "2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n",
        "3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 300 144] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>\nendobj\n",
        (
            f"4 0 obj\n<< /Length {len(stream.encode('utf-8'))} >>\nstream\n"
            f"{stream}\nendstream\nendobj\n"
        ),
        "5 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n",
    ]

    parts = ["%PDF-1.4\n"]
    offsets = [0]
    position = len(parts[0].encode("utf-8"))
    for obj in objects:
        offsets.append(position)
        parts.append(obj)
        position += len(obj.encode("utf-8"))

    xref = ["xref\n0 6\n", "0000000000 65535 f \n"]
    for offset in offsets[1:]:
        xref.append(f"{offset:010d} 00000 n \n")
    startxref = position
    trailer = f"trailer\n<< /Root 1 0 R /Size 6 >>\nstartxref\n{startxref}\n%%EOF\n"
    path.write_bytes("".join(parts + xref + [trailer]).encode("utf-8"))


def test_extract_text_document(tmp_path: Path) -> None:
    document_path = tmp_path / "notes.txt"
    document_path.write_text("First paragraph.\n\nSecond paragraph.", encoding="utf-8")

    result = extract_document(document_path)

    assert result.text.startswith("First paragraph")
    assert len(result.pages) == 1
    assert result.pages[0].page_number == 1


def test_extract_docx_document(tmp_path: Path) -> None:
    document_path = tmp_path / "notes.docx"
    doc = DocxDocument()
    doc.add_heading("Section 1", level=1)
    doc.add_paragraph("Body paragraph for the section.")
    doc.save(document_path)

    result = extract_document(document_path)

    assert "Body paragraph" in result.text
    assert result.pages[0].section == "Section 1"


def test_extract_pdf_document(tmp_path: Path) -> None:
    document_path = tmp_path / "notes.pdf"
    _write_simple_pdf(document_path, "PDF extraction text")

    result = extract_document(document_path)

    assert "PDF extraction text" in result.text
    assert result.pages[0].page_number == 1


def test_chunking_preserves_page_metadata() -> None:
    pages = [
        ExtractedPage(
            page_number=3,
            text="Alpha paragraph.\n\nBeta paragraph.",
            section="Overview",
        )
    ]

    chunks = chunk_pages(pages, source_label="sample.txt", max_chars=40)

    assert len(chunks) == 1
    assert chunks[0].page_number == 3
    assert chunks[0].section == "Overview"
    assert chunks[0].source_label == "sample.txt"
